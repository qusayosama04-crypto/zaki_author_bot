import os
import time
import telebot
import docx
import re
from google import genai
from flask import Flask
from threading import Thread

# 1. جلب المفاتيح من إعدادات السيرفر
BOT_TOKEN = os.getenv('BOT_TOKEN')
API_KEY = os.getenv('API_KEY')

bot = telebot.TeleBot(BOT_TOKEN)
client = genai.Client(api_key=API_KEY)

# دالة تنظيف النصوص من علامات الماركداون
def clean_markdown(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text) 
    text = re.sub(r'\*(.*?)\*', r'\1', text)     
    text = re.sub(r'#(.*?)\n', r'\1\n', text)    
    return text

@bot.message_handler(commands=['start'])
def send_welcome(message):
    bot.reply_to(message, "مرحباً يا مدير! 🫡\nأرسل لي أي موضوع تريد تأليف مجلد ضخم عنه، وسأقوم بالمهمة فوراً.")

@bot.message_handler(func=lambda message: True)
def write_epic_book(message):
    topic = message.text
    msg = bot.reply_to(message, f"⏳ جاري التخطيط لمجلد ضخم عن: {topic}\nأقوم الآن برسم الفهرس وتحديد الفصول... 🧠")

    try:
        # 2. طلب فهرس من 12 فصلاً
        index_prompt = f"أنت مؤلف كتب عالمي. اكتب لي فهرساً يتكون من 12 فصلاً لكتاب احترافي وشامل عن '{topic}'. أريد أسماء الفصول فقط مفصولة برمز النجمة (*) بدون أي أرقام أو نصوص أخرى."

        index_response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=index_prompt
        )

        # استخراج الفصول وتجاهل الفراغات
        chapters = [c.strip() for c in index_response.text.split('*') if len(c.strip()) > 3]

        if not chapters:
            bot.edit_message_text(chat_id=message.chat.id, message_id=msg.message_id, text="❌ حدث خطأ في استخراج الفهرس، يرجى المحاولة بعنوان آخر.")
            return

        # 3. تجهيز الوورد
        doc = docx.Document()
        doc.add_heading(f"كتاب: {topic}", 0)

        # 4. رحلة التأليف الطويلة (حلقة تكرارية)
        for i, chapter_title in enumerate(chapters, 1):
            bot.edit_message_text(chat_id=message.chat.id, message_id=msg.message_id, text=f"✍️ جاري تأليف الفصل {i} من {len(chapters)}...\n(العنوان: {chapter_title})\n\nالجودة العالية تحتاج وقتاً، احتسِ قهوتك ☕...")

            # هندسة الأوامر لكتابة فصل شديد الطول
            chapter_prompt = f"""
            أنت مؤلف كتب محترف وخبير عالمي. اكتب محتوى طويلاً جداً ومفصلاً وشاملاً (أكثر من 1500 كلمة) للفصل التالي: "{chapter_title}"
            وهو جزء من كتاب بعنوان: "{topic}".
            """
            chapter_response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=chapter_prompt
            )

            # تنظيف وإضافة الفصل
            clean_text = clean_markdown(chapter_response.text)
            doc.add_heading(chapter_title, level=1)
            doc.add_paragraph(clean_text)
            doc.add_page_break()

            # استراحة 15 ثانية للهروب من حظر جوجل وإعطاء الذكاء الاصطناعي فرصة للتنفس
            time.sleep(15)

        # 5. حفظ وإرسال
        safe_topic = topic.replace(' ', '_').replace('/', '_')
        file_name = f"{safe_topic}_Epic.docx"
        doc.save(file_name)
        
        with open(file_name, 'rb') as doc_file:
            bot.send_document(message.chat.id, doc_file)
        
        # تنظيف السيرفر من الملف بعد إرساله
        os.remove(file_name)
        bot.edit_message_text(chat_id=message.chat.id, message_id=msg.message_id, text="✅ تمت طباعة المجلد بنجاح وإرساله لك!")

    except Exception as e:
        print(f"Error: {e}")
        bot.edit_message_text(chat_id=message.chat.id, message_id=msg.message_id, text="❌ عذراً، حدث خطأ أثناء التأليف أو انقطع الاتصال. يرجى المحاولة لاحقاً.")

# --- بداية كود خدعة السيرفر (Flask) ---
app = Flask(__name__)

@app.route('/')
def home():
    return "🤖 المطبعة الملحمية تعمل 24/7 بنجاح! لا تنم يا سيرفر!"

def run():
    # هنا نجعل السيرفر يختار البورت تلقائياً لتجنب مشاكل Render
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)

def keep_alive():
    t = Thread(target=run)
    t.start()

# تشغيل خادم الويب الوهمي في الخلفية
keep_alive()
# --- نهاية كود الخدعة ---

print("🤖 المطبعة الملحمية تعمل الآن! جاهزة لكتابة المجلدات الضخمة...")
# تشغيل التيليجرام بوقت انتظار طويل لتجنب الانقطاع
bot.infinity_polling(timeout=90, long_polling_timeout=90)
